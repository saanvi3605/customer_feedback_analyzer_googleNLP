"""
Car Rental Feedback Analyzer (Watson AI)
Author: Saanvi Sharma
Repo: https://github.com/saanvi3605/customer_feedback_analyzer_googleNLP
"""

import streamlit as st
import pandas as pd
from docx import Document
from io import BytesIO
import plotly.express as px
import matplotlib.pyplot as plt
import re
from collections import defaultdict

from ibm_watson import NaturalLanguageUnderstandingV1
from ibm_cloud_sdk_core.authenticators import IAMAuthenticator
from ibm_watson.natural_language_understanding_v1 import Features, SentimentOptions

# ───────────────────────────── Streamlit page ─────────────────────────────
st.set_page_config(page_title="🚗 Car Rental Feedback Analyzer (Watson AI)", layout="wide")

# ───────────────────────────── Sidebar: Info & links ──────────────────────
st.sidebar.title("📘 Project Info")
st.sidebar.markdown(
    """
**Car Rental Feedback Analyzer**

Upload a CSV of customer reviews (or use demo data) and get an AI‑powered sentiment
breakdown plus automatic issue detection.  
Powered by **IBM Watson Natural Language Understanding**.
"""
)
st.sidebar.markdown("🔗 [View GitHub Repo](https://github.com/saanvi3605/customer_feedback_analyzer_googleNLP)")
st.sidebar.markdown("👨‍💻 Made by **Saanvi Sharma**")

# ───────────────────────────── Custom CSS ─────────────────────────────────
st.markdown(
    """
<style>
.main-header{
    text-align:center;padding:2rem 0;
    background:linear-gradient(135deg,#0062FF 0%,#0540A6 100%);
    border-radius:15px;margin-bottom:2rem;color:white}
.upload-area{border:2px dashed #0062FF;border-radius:15px;padding:2rem;text-align:center;
    background:#f0f5ff;margin-bottom:1rem}
.stat-card{background:white;padding:1rem;border-radius:12px;text-align:center;
    box-shadow:0 2px 10px rgba(0,0,0,0.1);margin-bottom:1rem}
</style>
""",
    unsafe_allow_html=True,
)

# ───────────────────────────── Demo data ──────────────────────────────────
def get_demo_data() -> pd.DataFrame:
    return pd.DataFrame(
        [
            dict(customer_id=4521, review="Car was clean but pickup took too long", rating=3),
            dict(customer_id=6723, review="Excellent service! Will rent again", rating=5),
            dict(customer_id=2891, review="Vehicle had maintenance issues", rating=2),
            dict(customer_id=3452, review="Great experience overall", rating=4),
            dict(customer_id=7834, review="Staff was rude and unhelpful", rating=1),
        ]
    )

# ───────────────────────────── IBM Watson init ───────────────────────────
@st.cache_resource
def init_watson():
    # 🔒 Replace these with your actual API key and service URL
    api_key = "2i7M1oiyhwXMHaeO1jOQkZdtHGjzCo_BreDibvSGn2cS"
    service_url = "https://api.au-syd.natural-language-understanding.watson.cloud.ibm.com/instances/fbf08a4e-d409-4965-894d-347f8faabce1"

    authenticator = IAMAuthenticator(api_key)
    nlu = NaturalLanguageUnderstandingV1(
        version="2023-06-01",
        authenticator=authenticator,
    )
    nlu.set_service_url(service_url)
    return nlu


# ───────────────────────────── Keyword map ───────────────────────────────
ISSUE_KEYWORDS = {
    "Service issue": r"(issue|problem|maintenance|break( |-)?down|engine|mechanic)",
    "Cleanliness concern": r"(dirty|unclean|smell|odor|stain|clean)",
    "Wait time issue": r"(slow|wait|delay|late|queue|pickup)",
    "Staff attitude": r"(rude|unhelpful|impolite|staff|customer service)",
    "Pricing complaint": r"(expensive|overpriced|costly|hidden fee|charge)",
    "Fuel / mileage": r"(fuel|gas|mileage|petrol|diesel)",
    "Insurance / docs": r"(insurance|document|paperwork)",
    "GPS / tech": r"(gps|navigation|bluetooth|usb|carplay)",
    "Child seat": r"(child seat|baby seat|booster)",
    "Damage": r"(scratch|dent|damage)",
}

# ───────────────────────────── Watson sentiment (cached) ──────────────────
@st.cache_data(show_spinner=False)
def watson_sentiment(text: str):
    nlu = init_watson()
    resp = nlu.analyze(text=text, features=Features(sentiment=SentimentOptions())).get_result()
    label = resp["sentiment"]["document"]["label"].lower()
    score = resp["sentiment"]["document"]["score"]
    return label, score

def analyze_review(text: str):
    sentiment, _ = watson_sentiment(text)
    issues = [label for label, pattern in ISSUE_KEYWORDS.items() if re.search(pattern, text, re.I)]
    return sentiment, sorted(set(issues))

# ───────────────────────────── Word report ───────────────────────────────
def generate_word_report(results):
    doc = Document()
    doc.add_heading("Car Rental Feedback Analysis Report (Watson AI)", 0)

    # summary numbers
    sentiment_counts = pd.Series([r["sentiment"] for r in results]).value_counts()
    doc.add_paragraph(f"Total Reviews: {len(results)}")
    doc.add_paragraph(f"Positive: {sentiment_counts.get('positive', 0)}")
    doc.add_paragraph(f"Neutral:  {sentiment_counts.get('neutral', 0)}")
    doc.add_paragraph(f"Negative: {sentiment_counts.get('negative', 0)}")

    # insert bar chart
    img_buf = BytesIO()
    plt.figure(figsize=(4, 3))
    plt.barh(sentiment_counts.index, sentiment_counts.values)
    plt.title("Sentiment Distribution")
    plt.xlabel("Count")
    plt.tight_layout()
    plt.savefig(img_buf, format="png")
    plt.close()
    img_buf.seek(0)
    doc.add_picture(img_buf, width=doc.sections[0].page_width * 0.55)

    # detailed table
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Review"
    hdr[1].text = "Sentiment"
    hdr[2].text = "Issues"

    for entry in results:
        row = table.add_row().cells
        row[0].text = entry["review"][:100]
        row[1].text = entry["sentiment"].title()
        row[2].text = ", ".join(entry["issues"]) or "None"

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

# ───────────────────────────── UI Header ─────────────────────────────────
st.markdown(
    """
<div class="main-header">
  <h1>🚗 Car Rental Feedback Analyzer (Watson AI)</h1>
  <p>Powered by IBM Watson Natural Language Understanding</p>
</div>
""",
    unsafe_allow_html=True,
)

# ───────────────────────────── File upload / demo ────────────────────────
st.markdown("## 📤 Upload Feedback CSV")
uploaded_file = st.file_uploader("Choose a CSV file", type=["csv"])
use_demo = st.checkbox("Use demo data instead")

df = pd.read_csv(uploaded_file) if uploaded_file else (get_demo_data() if use_demo else None)

if df is not None:
    st.write("Preview:")
    st.dataframe(df.head())

    if st.button("🔍 Start Analysis", type="primary"):
        with st.spinner("Analyzing reviews…"):
            results = []
            issue_counter = defaultdict(int)
            progress = st.progress(0)

            for i, row in df.iterrows():
                sentiment, issues = analyze_review(row["review"])
                results.append(
                    dict(
                        review=row["review"],
                        sentiment=sentiment,
                        issues=issues,
                        rating=row.get("rating", "N/A"),
                    )
                )
                for iss in issues:
                    issue_counter[iss] += 1
                progress.progress((i + 1) / len(df))
            progress.empty()

        # ───────────────────────── Summary output ────────────────────────
        st.info(f"🧾 Total Reviews Analyzed: {len(results)}")
        st.success("Analysis complete!")

        st.markdown("### 📊 Sentiment Summary")
        sentiment_counts = pd.Series([r["sentiment"] for r in results]).value_counts()

        c1, c2, c3 = st.columns(3)
        c1.metric("Positive", int(sentiment_counts.get("positive", 0)))
        c2.metric("Neutral",  int(sentiment_counts.get("neutral", 0)))
        c3.metric("Negative", int(sentiment_counts.get("negative", 0)))

        # Pie chart
        pie_fig = px.pie(
            sentiment_counts,
            values=sentiment_counts.values,
            names=sentiment_counts.index,
            title="Sentiment Distribution (Pie)",
        )
        st.plotly_chart(pie_fig, use_container_width=True)

        # Bar chart
        bar_fig = px.bar(
            sentiment_counts,
            x=sentiment_counts.index,
            y=sentiment_counts.values,
            text=sentiment_counts.values,
            title="Sentiment Distribution (Bar)",
            labels={"x": "Sentiment", "y": "Count"},
        )
        bar_fig.update_layout(yaxis=dict(dtick=1))
        st.plotly_chart(bar_fig, use_container_width=True)

        # Top issues
        if issue_counter:
            st.markdown("### 🚩 Top Issues")
            issues_df = (
                pd.DataFrame({"Issue": list(issue_counter.keys()), "Count": list(issue_counter.values())})
                .sort_values("Count", ascending=False)
            )
            st.dataframe(issues_df.head(10), use_container_width=True)

        # Word report
        st.markdown("### 📥 Download Word Report")
        docx_file = generate_word_report(results)
        st.download_button(
            label="📄 Download Report",
            data=docx_file,
            file_name="car_feedback_analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

# ───────────────────────────── Footer ────────────────────────────────────
st.markdown("---")
st.markdown(
    "<div style='text-align:center; color:gray'>Made by <strong>Saanvi Sharma</strong> • "
    "<a href='https://github.com/saanvi3605/customer_feedback_analyzer_googleNLP' target='_blank'>GitHub Repo</a>"
    "</div>",
    unsafe_allow_html=True,
)
